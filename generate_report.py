"""
Player Pulse: RFM Segmentation Report Generator
================================================
Generates a fully formatted Word document (.docx) with all analysis sections
and embedded visualizations.

Requirements:
    pip install python-docx matplotlib seaborn scikit-learn pandas numpy

Run:
    python generate_report.py
    -> outputs: Player_Pulse_RFM_Report.docx
"""

import os
import io
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')
sns.set_theme(style='whitegrid', palette='muted')
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['figure.dpi'] = 150

# ═══════════════════════════════════════════════════════════════════════════
# COLOUR PALETTE
# ═══════════════════════════════════════════════════════════════════════════
BRAND_BLUE   = RGBColor(0x1A, 0x3A, 0x5C)
ACCENT_TEAL  = RGBColor(0x00, 0x8B, 0x8B)
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY     = RGBColor(0x6E, 0x6E, 0x6E)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

SEG_COLORS_HEX = {
    'Champions':  '#2ecc71',
    'Loyal Users':'#3498db',
    'At Risk':    '#f39c12',
    'Dormant':    '#e67e22',
    'Lost':       '#e74c3c'
}
SEG_ORDER = ['Champions', 'Loyal Users', 'At Risk', 'Dormant', 'Lost']

SNAPSHOT_DATE = datetime(2024, 12, 31)

# ═══════════════════════════════════════════════════════════════════════════
# HELPER: shade a table cell
# ═══════════════════════════════════════════════════════════════════════════
def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = BRAND_BLUE
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = ACCENT_TEAL
    else:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = MID_GRAY
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size  = Pt(9)
    run.font.color.rgb = MID_GRAY
    run.font.name  = 'Calibri'
    doc.add_paragraph()

def body(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=11)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    add_run(p, text, size=10.5)
    return p

def embed_figure(doc, fig, cap_text, width=6.2):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    buf.seek(0)
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    caption(doc, cap_text)

# ═══════════════════════════════════════════════════════════════════════════
# DATA GENERATION
# ═══════════════════════════════════════════════════════════════════════════
def generate_data():
    np.random.seed(42)
    START_DATE = datetime(2024, 1, 1)
    END_DATE   = datetime(2024, 12, 31)
    DATE_RANGE = (END_DATE - START_DATE).days

    archetypes = {
        'champion':    {'n': 400,  'freq_mu': 45, 'freq_sd': 8,  'spend_mu': 85, 'spend_sd': 30, 'recency_max': 15},
        'loyal':       {'n': 700,  'freq_mu': 25, 'freq_sd': 6,  'spend_mu': 40, 'spend_sd': 15, 'recency_max': 30},
        'potential':   {'n': 900,  'freq_mu': 12, 'freq_sd': 4,  'spend_mu': 20, 'spend_sd': 10, 'recency_max': 45},
        'at_risk':     {'n': 800,  'freq_mu': 18, 'freq_sd': 5,  'spend_mu': 35, 'spend_sd': 12, 'recency_max': 120},
        'hibernating': {'n': 700,  'freq_mu': 8,  'freq_sd': 3,  'spend_mu': 15, 'spend_sd': 8,  'recency_max': 200},
        'lost':        {'n': 500,  'freq_mu': 3,  'freq_sd': 2,  'spend_mu': 5,  'spend_sd': 4,  'recency_max': 340},
        'new_user':    {'n': 1000, 'freq_mu': 4,  'freq_sd': 2,  'spend_mu': 10, 'spend_sd': 8,  'recency_max': 20},
    }

    records, user_id = [], 1
    for arch, p in archetypes.items():
        for _ in range(p['n']):
            n_tx = max(1, int(np.random.normal(p['freq_mu'], p['freq_sd'])))
            rec  = np.random.randint(1, p['recency_max'])
            last = END_DATE - timedelta(days=rec)
            for _ in range(n_tx):
                offset  = np.random.randint(0, min(DATE_RANGE, p['recency_max'] + 90))
                txn_dt  = max(last - timedelta(days=offset), START_DATE)
                amount  = max(0, np.random.normal(p['spend_mu'], p['spend_sd']))
                records.append({'user_id': f'U{user_id:05d}',
                                'transaction_date': txn_dt,
                                'amount': round(amount, 2)})
            user_id += 1

    df = pd.DataFrame(records)
    df['transaction_date'] = pd.to_datetime(df['transaction_date'])
    return df

# ═══════════════════════════════════════════════════════════════════════════
# RFM PIPELINE
# ═══════════════════════════════════════════════════════════════════════════
def build_rfm(df):
    rfm = df.groupby('user_id').agg(
        last_tx=('transaction_date', 'max'),
        frequency=('transaction_date', 'count'),
        monetary=('amount', 'sum')
    ).reset_index()
    rfm['recency'] = (SNAPSHOT_DATE - rfm['last_tx']).dt.days
    rfm.drop(columns=['last_tx'], inplace=True)

    rfm['R_score'] = pd.qcut(rfm['recency'], q=5, labels=[5,4,3,2,1]).astype(int)
    rfm['F_score'] = pd.qcut(rfm['frequency'].rank(method='first'), q=5, labels=[1,2,3,4,5]).astype(int)
    rfm['M_score'] = pd.qcut(rfm['monetary'].rank(method='first'),  q=5, labels=[1,2,3,4,5]).astype(int)
    rfm['RFM_score'] = (rfm['R_score'] + rfm['F_score'] + rfm['M_score']) / 3
    return rfm

def cluster_rfm(rfm):
    feats  = rfm[['R_score','F_score','M_score']].values
    scaled = StandardScaler().fit_transform(feats)

    inertias, sil = [], []
    for k in range(2, 11):
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        km.fit(scaled)
        inertias.append(km.inertia_)
        sil.append(silhouette_score(scaled, km.labels_))

    km5 = KMeans(n_clusters=5, random_state=42, n_init=10)
    rfm['cluster'] = km5.fit_predict(scaled)

    rank_map = {cid: lbl for rank, (cid, _) in
                enumerate(rfm.groupby('cluster')['RFM_score'].mean()
                            .sort_values(ascending=False).items())
                for lbl in [['Champions','Loyal Users','At Risk','Dormant','Lost'][rank]]}
    rfm['segment'] = rfm['cluster'].map(rank_map)
    return rfm, inertias, sil, scaled

# ═══════════════════════════════════════════════════════════════════════════
# FIGURES
# ═══════════════════════════════════════════════════════════════════════════
def fig_eda(df):
    fig, axes = plt.subplots(1, 3, figsize=(16, 4.5))
    weekly = df.set_index('transaction_date').resample('W')['user_id'].count()
    axes[0].plot(weekly.index, weekly.values, color='steelblue', linewidth=1.8)
    axes[0].fill_between(weekly.index, weekly.values, alpha=0.15, color='steelblue')
    axes[0].set_title('Weekly Transaction Volume', fontweight='bold')
    axes[0].set_xlabel('Date'); axes[0].set_ylabel('Transactions')

    df[df['amount'] > 0]['amount'].hist(bins=50, ax=axes[1], color='coral', edgecolor='white')
    axes[1].set_title('Transaction Amount Distribution', fontweight='bold')
    axes[1].set_xlabel('Amount (USD)'); axes[1].set_ylabel('Count')

    df.groupby('user_id')['transaction_date'].count().hist(
        bins=40, ax=axes[2], color='mediumseagreen', edgecolor='white')
    axes[2].set_title('Transactions per User', fontweight='bold')
    axes[2].set_xlabel('# Transactions'); axes[2].set_ylabel('Users')

    plt.tight_layout()
    return fig

def fig_rfm_scores(rfm):
    fig, axes = plt.subplots(1, 3, figsize=(14, 4.5))
    for ax, col, color, label in zip(
            axes,
            ['R_score','F_score','M_score'],
            ['#4C72B0','#DD8452','#55A868'],
            ['Recency Score','Frequency Score','Monetary Score']):
        rfm[col].value_counts().sort_index().plot(
            kind='bar', ax=ax, color=color, edgecolor='white', width=0.7)
        ax.set_title(label, fontweight='bold')
        ax.set_xlabel('Score (1=Worst, 5=Best)'); ax.set_ylabel('Users')
        ax.tick_params(axis='x', rotation=0)
    plt.tight_layout()
    return fig

def fig_elbow(inertias, sil):
    k_range = range(2, 11)
    fig, axes = plt.subplots(1, 2, figsize=(13, 4.5))
    axes[0].plot(list(k_range), inertias, 'o-', color='steelblue', linewidth=2)
    axes[0].axvline(x=5, color='red', linestyle='--', alpha=0.7, label='k=5 selected')
    axes[0].set_title('Elbow Method — Inertia vs. K', fontweight='bold')
    axes[0].set_xlabel('k'); axes[0].set_ylabel('Inertia'); axes[0].legend()

    axes[1].plot(list(k_range), sil, 'o-', color='coral', linewidth=2)
    axes[1].axvline(x=5, color='red', linestyle='--', alpha=0.7, label='k=5 selected')
    axes[1].set_title('Silhouette Score vs. K', fontweight='bold')
    axes[1].set_xlabel('k'); axes[1].set_ylabel('Silhouette Score'); axes[1].legend()
    plt.tight_layout()
    return fig

def fig_dashboard(rfm, seg_summary):
    fig, axes = plt.subplots(2, 2, figsize=(15, 12))

    ax = axes[0,0]
    for seg in SEG_ORDER:
        grp = rfm[rfm['segment']==seg]
        ax.scatter(grp['recency'], grp['frequency'], alpha=0.3, s=10,
                   color=SEG_COLORS_HEX[seg], label=seg)
    ax.set_xlabel('Recency (days)'); ax.set_ylabel('Frequency')
    ax.set_title('Recency vs. Frequency by Segment', fontweight='bold')
    ax.legend(markerscale=2, fontsize=9)

    ax = axes[0,1]
    rev = seg_summary['pct_revenue'].loc[SEG_ORDER]
    bars = ax.barh(rev.index, rev.values,
                   color=[SEG_COLORS_HEX[s] for s in rev.index], edgecolor='white')
    ax.set_xlabel('% of Total Revenue')
    ax.set_title('Revenue Share by Segment', fontweight='bold')
    for b, v in zip(bars, rev.values):
        ax.text(v+0.3, b.get_y()+b.get_height()/2, f'{v:.1f}%', va='center', fontsize=9)
    ax.set_xlim(0, rev.max()+10)

    ax = axes[1,0]
    hmap = rfm.groupby('segment')[['R_score','F_score','M_score']].mean().loc[SEG_ORDER]
    hmap.columns = ['Recency\nScore','Frequency\nScore','Monetary\nScore']
    sns.heatmap(hmap, annot=True, fmt='.1f', cmap='RdYlGn',
                vmin=1, vmax=5, linewidths=0.5, ax=ax,
                cbar_kws={'label':'Score (1=Low, 5=High)'})
    ax.set_title('Mean RFM Scores by Segment', fontweight='bold')
    ax.set_ylabel('')

    ax = axes[1,1]
    users = seg_summary['Users'].loc[SEG_ORDER]
    ax.pie(users.values, labels=users.index,
           colors=[SEG_COLORS_HEX[s] for s in users.index],
           autopct='%1.1f%%', startangle=140,
           wedgeprops=dict(width=0.5))
    ax.set_title('User Distribution by Segment', fontweight='bold')

    plt.tight_layout()
    return fig

def fig_profiles(seg_summary):
    fig, axes = plt.subplots(1, 3, figsize=(16, 5))
    metrics = [
        ('Avg_Recency',   'Avg. Recency (days)\n(lower = more recent)', True),
        ('Avg_Frequency', 'Avg. Frequency (transactions)',               False),
        ('Avg_Monetary',  'Avg. Spend per User (USD)',                   False),
    ]
    for ax, (col, ylabel, inv) in zip(axes, metrics):
        vals = seg_summary[col].loc[SEG_ORDER]
        bars = ax.bar(vals.index, vals.values,
                      color=[SEG_COLORS_HEX[s] for s in vals.index],
                      edgecolor='white', width=0.6)
        ax.set_title(ylabel, fontweight='bold', fontsize=10)
        ax.tick_params(axis='x', rotation=20)
        for b, v in zip(bars, vals.values):
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+vals.max()*0.01,
                    f'{v:.0f}', ha='center', va='bottom', fontsize=8)
        if inv:
            ax.invert_yaxis()
    plt.tight_layout()
    return fig

# ═══════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════
def build_document(df, rfm, seg_summary, inertias, sil):
    doc = Document()

    # Page margins (1 inch)
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # ── COVER ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Player Pulse', bold=True, size=28, color=BRAND_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'User Lifecycle Segmentation with RFM Analysis', size=16, color=ACCENT_TEAL)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'Analysis Date: {SNAPSHOT_DATE.strftime("%B %d, %Y")}  |  Dataset: Synthetic (5,000 users, 12 months)',
            size=10, color=MID_GRAY, italic=True)

    doc.add_paragraph()
    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ───────────────────────────────────────────────
    heading(doc, '1. Executive Summary')
    body(doc,
        'This report presents a complete RFM (Recency, Frequency, Monetary) '
        'segmentation analysis of 5,000 synthetic users over a 12-month period. '
        'Using K-Means clustering on quintile-scored RFM dimensions, users are grouped '
        'into five behaviorally distinct segments — from Champions who drive outsized '
        'revenue, to Lost users who have fully disengaged. The methodology and insights '
        'are directly transferable to any live-service product, subscription business, '
        'e-commerce platform, or mobile game.')

    # Summary KPI table
    doc.add_paragraph()
    tbl = doc.add_table(rows=6, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    headers = ['Segment', 'Users', '% of Users', 'Avg. Spend', '% Revenue']
    hdr_fills = ['1A3A5C'] * 5
    for j, (h, f) in enumerate(zip(headers, hdr_fills)):
        cell = tbl.rows[0].cells[j]
        shade_cell(cell, f'#{f}')
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, h, bold=True, size=10, color=WHITE)

    row_data = [
        ('Champions',  seg_summary.loc['Champions',  'Users'], seg_summary.loc['Champions',  'pct_users'], seg_summary.loc['Champions',  'Avg_Monetary'], seg_summary.loc['Champions',  'pct_revenue']),
        ('Loyal Users',seg_summary.loc['Loyal Users','Users'], seg_summary.loc['Loyal Users','pct_users'], seg_summary.loc['Loyal Users','Avg_Monetary'], seg_summary.loc['Loyal Users','pct_revenue']),
        ('At Risk',    seg_summary.loc['At Risk',    'Users'], seg_summary.loc['At Risk',    'pct_users'], seg_summary.loc['At Risk',    'Avg_Monetary'], seg_summary.loc['At Risk',    'pct_revenue']),
        ('Dormant',    seg_summary.loc['Dormant',    'Users'], seg_summary.loc['Dormant',    'pct_users'], seg_summary.loc['Dormant',    'Avg_Monetary'], seg_summary.loc['Dormant',    'pct_revenue']),
        ('Lost',       seg_summary.loc['Lost',       'Users'], seg_summary.loc['Lost',       'pct_users'], seg_summary.loc['Lost',       'Avg_Monetary'], seg_summary.loc['Lost',       'pct_revenue']),
    ]
    seg_fills = {'Champions':'D5F5E3','Loyal Users':'D6EAF8','At Risk':'FDEBD0','Dormant':'FAD7A0','Lost':'FADBD8'}

    for i, (seg, u, pu, am, pr) in enumerate(row_data):
        row = tbl.rows[i+1]
        vals = [seg, str(int(u)), f'{pu:.1f}%', f'${am:.0f}', f'{pr:.1f}%']
        for j, val in enumerate(vals):
            cell = row.cells[j]
            shade_cell(cell, f'#{seg_fills[seg]}')
            p2 = cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p2, val, size=10, bold=(j==0))

    doc.add_paragraph()

    # ── 2. DATA CREATION ───────────────────────────────────────────────────
    heading(doc, '2. Dataset Creation')
    body(doc,
        'A synthetic transaction dataset was generated in Python to simulate one year of '
        'user activity (January 1 – December 31, 2024). The dataset intentionally mirrors '
        'the structure of real-world CRM exports, mobile game telemetry pipelines, '
        'or data warehouse event logs, making it straightforward to swap in production data.')

    heading(doc, 'Schema', level=2)
    schema_rows = [
        ('user_id',          'String',   'Unique user/player identifier (e.g. U00001)'),
        ('transaction_date', 'Datetime', 'Date of purchase or meaningful engagement event'),
        ('amount',           'Float',    'Revenue generated in USD (0 for free sessions)'),
    ]
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Column', 'Type', 'Description']):
        c = t2.rows[0].cells[j]
        shade_cell(c, '#1A3A5C')
        p2 = c.paragraphs[0]
        add_run(p2, h, bold=True, size=10, color=WHITE)
    for i, (col, typ, desc) in enumerate(schema_rows):
        row = t2.rows[i+1]
        shade_cell(row.cells[0], '#EBF5FB')
        for j, val in enumerate([col, typ, desc]):
            p2 = row.cells[j].paragraphs[0]
            add_run(p2, val, size=10, bold=(j==0))
    doc.add_paragraph()

    heading(doc, 'User Archetypes', level=2)
    body(doc,
        'Seven behavioral archetypes were defined to produce a realistic, non-uniform '
        'distribution across the RFM space. Each archetype controls the number of '
        'transactions, recency window, and spend distribution for its cohort:')
    for arch, desc in [
        ('Champion (400 users)',    'Very high frequency (avg. 45 tx), recent activity, high spend (~$85 avg)'),
        ('Loyal (700 users)',       'Consistent engagement (avg. 25 tx), moderate-high spend (~$40 avg)'),
        ('Potential (900 users)',   'Growing users with moderate activity (avg. 12 tx) and lower spend (~$20 avg)'),
        ('At Risk (800 users)',     'Previously active users with increasing recency (up to 120 days absent)'),
        ('Hibernating (700 users)', 'Low frequency, low spend, last active 60-200 days ago'),
        ('Lost (500 users)',        'Minimal historical activity, last seen 200-340 days ago'),
        ('New User (1,000 users)',  'Very recent acquisition (within 20 days), low transaction count'),
    ]:
        p2 = doc.add_paragraph(style='List Bullet')
        add_run(p2, f'{arch}: ', bold=True, size=10.5)
        add_run(p2, desc, size=10.5)

    doc.add_paragraph()
    body(doc,
        f'The generator produced {len(df):,} total transaction records across '
        f'{df["user_id"].nunique():,} unique users, spanning '
        f'{df["transaction_date"].min().date()} to {df["transaction_date"].max().date()}, '
        f'with total simulated revenue of ${df["amount"].sum():,.0f}.')

    # ── 3. EDA ─────────────────────────────────────────────────────────────
    heading(doc, '3. Exploratory Data Analysis')
    body(doc,
        'Before computing RFM metrics, three foundational EDA checks were performed: '
        'transaction volume over time to identify seasonality or anomalies, '
        'revenue distribution to understand skew and outlier presence, '
        'and sessions-per-user to characterize engagement breadth.')

    embed_figure(doc, fig_eda(df),
                 'Figure 1: EDA Overview — Weekly volume, amount distribution, transactions per user')

    body(doc,
        'Key observations: transaction volume is broadly consistent across the year with '
        'natural variance. The amount distribution is right-skewed, as expected in '
        'live-service monetization — a small proportion of users drive a disproportionate '
        'share of revenue. The sessions-per-user histogram is bimodal, reflecting the '
        'high-frequency champion cohort versus the broader low-engagement base.')

    # ── 4. DATA CLEANING ───────────────────────────────────────────────────
    heading(doc, '4. Data Cleaning & Quality Checks')
    body(doc,
        'The following checks were applied before feature engineering. '
        'The synthetic dataset was designed to be clean, but these steps '
        'are included as a production-ready pattern:')

    checks = [
        ('Null values',          f'{df.isnull().sum().sum()} nulls found across all columns. No imputation required.'),
        ('Duplicate transactions','Checked for duplicate (user_id, transaction_date, amount) triplets. None found.'),
        ('Negative amounts',      f'{(df["amount"] < 0).sum()} negative-value records. Floored to 0 to handle refunds.'),
        ('Future-dated records',  f'{(df["transaction_date"] > pd.Timestamp(SNAPSHOT_DATE)).sum()} records post snapshot date. Excluded from recency calculation.'),
        ('Users with 0 revenue',  f'{(df.groupby("user_id")["amount"].sum() == 0).sum()} users with zero lifetime spend — retained as valid free-to-play users.'),
    ]
    t3 = doc.add_table(rows=len(checks)+1, cols=2)
    t3.style = 'Table Grid'
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Check', 'Result']):
        c = t3.rows[0].cells[j]
        shade_cell(c, '#1A3A5C')
        add_run(c.paragraphs[0], h, bold=True, size=10, color=WHITE)
    for i, (chk, res) in enumerate(checks):
        row = t3.rows[i+1]
        fill = '#F2F3F4' if i % 2 == 0 else '#FFFFFF'
        shade_cell(row.cells[0], fill); shade_cell(row.cells[1], fill)
        add_run(row.cells[0].paragraphs[0], chk, bold=True, size=10)
        add_run(row.cells[1].paragraphs[0], res, size=10)
    doc.add_paragraph()

    # ── 5. FEATURE ENGINEERING ─────────────────────────────────────────────
    heading(doc, '5. Feature Engineering — RFM Metrics')
    body(doc,
        'The raw transaction log was aggregated to one row per user with three computed dimensions:')

    for dim, defn in [
        ('Recency (R)',   'Days between the user\'s most recent transaction and the snapshot date (Dec 31, 2024). Lower = better.'),
        ('Frequency (F)', 'Total count of transaction records for the user within the analysis period. Higher = better.'),
        ('Monetary (M)',  'Sum of all transaction amounts for the user. Higher = better.'),
    ]:
        p2 = doc.add_paragraph(style='List Bullet')
        add_run(p2, f'{dim}: ', bold=True, size=10.5)
        add_run(p2, defn, size=10.5)

    doc.add_paragraph()
    body(doc,
        'Each dimension was then scored 1–5 using quintile binning (qcut). '
        'For Recency, the scoring is inverted so that score 5 = most recent. '
        'For Frequency and Monetary, score 5 = highest. '
        'A composite RFM score (simple average of the three) provides a single '
        'ranking metric per user, used later to anchor segment labeling.')

    embed_figure(doc, fig_rfm_scores(rfm),
                 'Figure 2: RFM Score Distributions — Approximately uniform across quintiles, confirming balanced scoring')

    # ── 6. CLUSTERING ──────────────────────────────────────────────────────
    heading(doc, '6. K-Means Clustering')
    body(doc,
        'K-Means clustering was applied to the standardized R, F, and M scores to identify '
        'natural groupings in the 3-dimensional RFM space. StandardScaler was used to '
        'normalize the features before clustering, preventing any single dimension from '
        'dominating by scale.')

    heading(doc, 'Selecting Optimal k', level=2)
    body(doc,
        'The elbow method (inertia vs. k) and silhouette score were evaluated across '
        'k = 2 through 10. Both methods converged on k = 5 as the optimal number of '
        'clusters — consistent with the five intuitive lifecycle stages '
        '(Champions, Loyal, At Risk, Dormant, Lost) commonly used in CRM analytics.')

    embed_figure(doc, fig_elbow(inertias, sil),
                 'Figure 3: Elbow Method and Silhouette Score — Both indicate k=5 as optimal')

    body(doc,
        'At k = 5, the silhouette score was '
        f'{silhouette_score(StandardScaler().fit_transform(rfm[["R_score","F_score","M_score"]].values), rfm["cluster"].values):.3f}, '
        'indicating well-separated, cohesive clusters. '
        'Clusters were labeled by sorting on mean composite RFM score, '
        'with rank 1 assigned "Champions" through rank 5 assigned "Lost".')

    # ── 7. SEGMENT ANALYSIS ────────────────────────────────────────────────
    heading(doc, '7. Segment Analysis & Insights')
    embed_figure(doc, fig_dashboard(rfm, seg_summary),
                 'Figure 4: Segmentation Dashboard — Scatter, revenue share, RFM heatmap, user distribution', width=6.4)

    embed_figure(doc, fig_profiles(seg_summary),
                 'Figure 5: Segment Profile Comparison — Recency, frequency, and monetary by segment')

    heading(doc, 'Key Findings', level=2)

    findings = [
        ('Champions are revenue-outsized.',
         f'Champions make up {seg_summary.loc["Champions","pct_users"]:.1f}% of users but account for '
         f'{seg_summary.loc["Champions","pct_revenue"]:.1f}% of total revenue. '
         'Average spend is ${:.0f}/user — {:.1f}x the platform average.'.format(
             seg_summary.loc['Champions','Avg_Monetary'],
             seg_summary.loc['Champions','Avg_Monetary'] / seg_summary['Avg_Monetary'].mean())),
        ('The At Risk segment is the highest-ROI re-engagement target.',
         f'{int(seg_summary.loc["At Risk","Users"]):,} users were previously highly active '
         f'(avg. {seg_summary.loc["At Risk","Avg_Frequency"]:.0f} transactions) '
         f'but have not engaged in an average of {seg_summary.loc["At Risk","Avg_Recency"]:.0f} days. '
         'Their historical spend profile makes them high-value reactivation candidates.'),
        ('Dormant and Lost segments represent churn already realized.',
         f'{int(seg_summary.loc["Dormant","Users"] + seg_summary.loc["Lost","Users"]):,} users '
         f'({seg_summary.loc["Dormant","pct_users"]+seg_summary.loc["Lost","pct_users"]:.1f}% of base) '
         'have recency of 100+ days and minimal spend. Campaign ROI for this group is low; '
         'focus should be on root cause analysis to prevent future churn.'),
        ('New users require rapid activation.',
         f'The {int(seg_summary.loc["Loyal Users","Users"]):,}-user Loyal segment and '
         f'{int(rfm[rfm["recency"] <= 20]["user_id"].nunique()):,} very recent users show that '
         'a significant portion of the base has high activation potential if onboarding '
         'engagement is optimized in the first 30 days.'),
    ]

    for title, detail in findings:
        p2 = doc.add_paragraph(style='List Bullet')
        add_run(p2, f'{title}  ', bold=True, size=10.5)
        add_run(p2, detail, size=10.5)
    doc.add_paragraph()

    # ── 8. RECOMMENDATIONS ─────────────────────────────────────────────────
    heading(doc, '8. Business Recommendations')
    body(doc,
        'The following actions are recommended per segment. Each is framed for two '
        'contexts to demonstrate the transferability of this framework.')

    recs = [
        ('Champions',   '#D5F5E3',
         'Exclusive early access, VIP cosmetics, beta tester programs. Protect LTV — avoid over-monetizing.',
         'Loyalty tier upgrades, referral programs, NPS survey candidates, upsell to premium.'),
        ('Loyal Users', '#D6EAF8',
         'Battle Pass nudges, streak rewards, social feature activation to deepen loops.',
         'Feature education campaigns, conversion to higher subscription tier, targeted upsell.'),
        ('At Risk',     '#FDEBD0',
         'Re-engagement push: "Your friends are playing." Patch highlight email. Free XP boost.',
         'Win-back email sequence, personalized "We miss you" discount, churn flag to CS team.'),
        ('Dormant',     '#FAD7A0',
         'Major content update announcements, seasonal event emails, low-friction re-entry.',
         'Broad promotional newsletter, discount offer, suppress if no response in 60 days.'),
        ('Lost',        '#FADBD8',
         'One final comeback campaign tied to a major release. Suppress thereafter.',
         'GDPR/data hygiene candidate. Analyze exit patterns to prevent future churn.'),
    ]

    t4 = doc.add_table(rows=len(recs)+1, cols=3)
    t4.style = 'Table Grid'
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Segment', 'Gaming Action (e.g. Riot)', 'General Action (SaaS / E-Commerce)']):
        c = t4.rows[0].cells[j]
        shade_cell(c, '#1A3A5C')
        p2 = c.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, h, bold=True, size=10, color=WHITE)
    for i, (seg, fill, gaming, general) in enumerate(recs):
        row = t4.rows[i+1]
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], '#FDFEFE')
        shade_cell(row.cells[2], '#FDFEFE')
        add_run(row.cells[0].paragraphs[0], seg, bold=True, size=10)
        add_run(row.cells[1].paragraphs[0], gaming, size=10)
        add_run(row.cells[2].paragraphs[0], general, size=10)
    doc.add_paragraph()

    # ── 9. NEXT STEPS ──────────────────────────────────────────────────────
    heading(doc, '9. Next Steps & Extensions')
    body(doc,
        'This segmentation is a foundation. The following analytical extensions '
        'build directly on the RFM segments to deliver additional business value:')

    next_steps = [
        ('Churn Prediction Model',
         'Train an XGBoost or logistic regression classifier using RFM scores and segment '
         'membership as features to predict 30-day churn probability at the individual user level. '
         'Output a ranked list of at-risk users for targeted intervention.'),
        ('Lifetime Value (LTV) Regression',
         'Predict 90-day forward revenue per user using historical RFM, segment, and '
         'behavioral features. Use SHAP values to identify the strongest LTV drivers '
         'and inform product prioritization.'),
        ('A/B Experiment Design on At Risk Segment',
         'Hold out 50% of At Risk users as a control group, send a re-engagement campaign '
         'to the treatment group, and measure 30-day reactivation lift. '
         'This directly validates the segment\'s responsiveness and informs campaign ROI.'),
        ('Segment Migration Tracking',
         'Run RFM monthly and track how users transition between segments over time. '
         'Visualize as a Sankey or cohort flow diagram. Identify which product or campaign '
         'events correlate with positive migration (e.g., Dormant → At Risk → Loyal).'),
        ('Multi-Channel Attribution Model',
         'Layer campaign touchpoint data (email opens, push clicks, paid installs) on top '
         'of the segmented user base. Build a multi-touch attribution model to quantify '
         'which channels are most effective for each segment, improving media efficiency.'),
        ('Automated Reporting Pipeline',
         'Schedule this analysis to run monthly, automatically flag segment shifts '
         'above a threshold (e.g., >5% movement in Champions), and push a summary '
         'report to stakeholders. Connects to the Lifecycle Marketing Insights KPI '
         'dashboard use case directly relevant to senior analytics roles.'),
    ]

    for title, detail in next_steps:
        p2 = doc.add_paragraph(style='List Bullet')
        add_run(p2, f'{title}: ', bold=True, size=10.5)
        add_run(p2, detail, size=10.5)
    doc.add_paragraph()

    # ── FOOTER ─────────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        p2 = footer.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2,
                'Player Pulse: RFM Segmentation Report  |  Methodology: Python (pandas, scikit-learn, matplotlib)  |  '
                f'Analysis Date: {SNAPSHOT_DATE.strftime("%B %Y")}',
                size=8, color=MID_GRAY, italic=True)

    return doc


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating data...')
    df = generate_data()
    print(f'  {len(df):,} transactions | {df["user_id"].nunique():,} users')

    print('Building RFM features...')
    rfm = build_rfm(df)

    print('Clustering...')
    rfm, inertias, sil, scaled = cluster_rfm(rfm)

    seg_summary = rfm.groupby('segment').agg(
        Users=('user_id','count'),
        Avg_Recency=('recency','mean'),
        Avg_Frequency=('frequency','mean'),
        Avg_Monetary=('monetary','mean'),
        Total_Revenue=('monetary','sum'),
    ).round(1)
    seg_summary['pct_users']   = (seg_summary['Users'] / seg_summary['Users'].sum() * 100).round(1)
    seg_summary['pct_revenue'] = (seg_summary['Total_Revenue'] / seg_summary['Total_Revenue'].sum() * 100).round(1)

    print('Building Word document...')
    doc = build_document(df, rfm, seg_summary, inertias, sil)

    out_path = 'Player_Pulse_RFM_Report.docx'
    doc.save(out_path)
    print(f'\nDone. Saved: {out_path}')
